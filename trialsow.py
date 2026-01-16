import streamlit as st
from datetime import date
import io
import re
import os
import pandas as pd
import requests

# --- FILE PATHS ---
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
SOW_DIAGRAM_MAP = {
    "L1 Support Bot POC SOW": os.path.join(BASE_DIR, "diagrams", "L1 Support Bot POC SOW.png"),
    "Ready Search POC Scope of Work Document": os.path.join(BASE_DIR, "diagrams", "Ready Search POC Scope of Work Document.png"),
    "AI based Image Enhancement POC SOW": os.path.join(BASE_DIR, "diagrams", "AI based Image Enhancement POC SOW.png"),
    "Beauty Advisor POC SOW": os.path.join(BASE_DIR, "diagrams", "Beauty Advisor POC SOW.png"),
    "AI based Image Inspection POC SOW": os.path.join(BASE_DIR, "diagrams", "AI based Image Inspection POC SOW.png"),
    "Gen AI for SOP POC SOW": os.path.join(BASE_DIR, "diagrams", "Gen AI for SOP POC SOW.png"),
    "Project Scope Document": os.path.join(BASE_DIR, "diagrams", "Project Scope Document.png"),
    "Gen AI Speech To Speech": os.path.join(BASE_DIR, "diagrams", "Gen AI Speech To Speech.png"),
    "PoC Scope Document": os.path.join(BASE_DIR, "diagrams", "PoC Scope Document.png")
}

SOW_COST_TABLE_MAP = {
    "L1 Support Bot POC SOW": {"POC Cost": "3,536.40 USD", "AWS Calculator Link": "https://calculator.aws/#/"},
    "Beauty Advisor POC SOW": {"POC Cost": "4,725.66 USD", "Prod Cost": "5,701.48 USD", "AWS Calculator Link": "https://calculator.aws/#/"},
    "Ready Search POC Scope of Work Document": {"POC Cost": "2,641.40 USD", "AWS Calculator Link": "https://calculator.aws/#/"},
    "AI based Image Enhancement POC SOW": {"POC Cost": "2,814.34 USD", "AWS Calculator Link": "https://calculator.aws/#/"},
    "AI based Image Inspection POC SOW": {"POC Cost": "3,536.40 USD", "AWS Calculator Link": "https://calculator.aws/#/"},
    "Gen AI for SOP POC SOW": {"POC Cost": "2,110.30 USD", "AWS Calculator Link": "https://calculator.aws/#/"},
    "Project Scope Document": {"Prod Cost": "2,993.60 USD", "AWS Calculator Link": "https://calculator.aws/#/"},
    "Gen AI Speech To Speech": {"Prod Cost": "2,124.23 USD", "AWS Calculator Link": "https://calculator.aws/#/"},
    "PoC Scope Document": {"Amazon Bedrock": "1,000 USD", "Total": "$ 3,150", "AWS Calculator Link": "https://calculator.aws/#/"}
}

# --- STREAMLIT CONFIG ---
st.set_page_config(page_title="GenAI SOW Architect", layout="wide", page_icon="📄", initial_sidebar_state="expanded")

# --- CUSTOM CSS ---
st.markdown("""
<style>
.main { background-color: #f8fafc; }
.stButton>button { border-radius: 8px; font-weight: 600; }
.stTextArea textarea { border-radius: 10px; }
.stTextInput input { border-radius: 8px; }
.block-container { padding-top: 1.5rem; }
.sow-preview {
    background-color: white;
    padding: 40px;
    border-radius: 12px;
    border: 1px solid #e2e8f0;
    font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
    line-height: 1.7;
    color: #1e293b;
    box-shadow: 0 4px 6px -1px rgb(0 0 0 / 0.1);
}
h1, h2, h3 { color: #0f172a; }
.stTabs [data-baseweb="tab-list"] { gap: 24px; }
.stTabs [data-baseweb="tab"] { height: 50px; white-space: pre-wrap; font-weight: 600; }
[data-testid="stExpander"] { border: none; box-shadow: none; background: transparent; }
.stakeholder-header { 
    background-color: #f1f5f9; 
    padding: 8px 12px; 
    border-radius: 6px; 
    margin-bottom: 10px; 
    font-weight: bold;
    color: #334155;
    border-left: 4px solid #3b82f6;
}
</style>
""", unsafe_allow_html=True)

# --- DOCX GENERATION FUNCTION ---
def create_docx_logic(text_content, branding_info, sow_type_name):
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Cover Page
    if branding_info.get('aws_pn_logo_bytes'):
        p_top = doc.add_paragraph()
        p_top.alignment = WD_ALIGN_PARAGRAPH.LEFT
        try:
            p_top.add_run().add_picture(io.BytesIO(branding_info['aws_pn_logo_bytes']), width=Inches(1.0))
        except:
            p_top.add_run("AWS PN").bold = True

    doc.add_paragraph("\n" * 3)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(branding_info['sow_name'])
    run.font.size = Pt(26)
    run.font.bold = True

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle_p.add_run("Scope of Work Document")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_paragraph("\n" * 4)

    # Logos table (3 fixed + optional customer)
    logo_table = doc.add_table(rows=1, cols=3)
    logo_table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def insert_logo(cell, bytes_data, width_val, fallback_text):
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.text = ""
        if bytes_data:
            try:
                p.add_run().add_picture(io.BytesIO(bytes_data), width=Inches(width_val))
            except:
                p.add_run(fallback_text).bold = True
        else:
            p.add_run(fallback_text).bold = True

    insert_logo(logo_table.rows[0].cells[0], branding_info.get('customer_logo_bytes'), 1.4, "[Customer Logo]")
    insert_logo(logo_table.rows[0].cells[1], branding_info.get('oneture_logo_bytes'), 2.2, "ONETURE")
    insert_logo(logo_table.rows[0].cells[2], branding_info.get('aws_adv_logo_bytes'), 1.3, "AWS Advanced")

    doc.add_paragraph("\n" * 4)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run(branding_info['doc_date_str'])
    run.font.size = Pt(12)
    run.font.bold = True

    doc.add_page_break()

    # Content
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    lines = text_content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue
        clean_text = re.sub(r'\*+|^#+\s*', '', line).strip()
        upper_text = clean_text.upper()

        # Section 4: Architecture + Cost Table
        if "4 SOLUTION ARCHITECTURE" in upper_text:
            doc.add_heading(clean_text, level=1)
            diagram_path = SOW_DIAGRAM_MAP.get(sow_type_name)
            if diagram_path and os.path.exists(diagram_path):
                try:
                    doc.add_picture(diagram_path, width=Inches(6.0))
                    p_cap = doc.add_paragraph(f"{sow_type_name} – Architecture Diagram")
                    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except:
                    doc.add_paragraph("[Architecture Diagram Missing]")

            cost_info = SOW_COST_TABLE_MAP.get(sow_type_name)
            if cost_info:
                table = doc.add_table(rows=1, cols=len(cost_info))
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                for idx, key in enumerate(cost_info.keys()):
                    hdr_cells[idx].text = key
                row_cells = table.add_row().cells
                for idx, key in enumerate(cost_info.keys()):
                    row_cells[idx].text = str(cost_info[key])
            i += 1
            continue

        if line.startswith('# '):
            doc.add_heading(clean_text, level=1)
        elif line.startswith('## '):
            doc.add_heading(clean_text, level=2)
        elif line.startswith('### '):
            doc.add_heading(clean_text, level=3)
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(clean_text[2:], style='List Bullet')
        else:
            doc.add_paragraph(clean_text)
        i += 1

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- SESSION STATE ---
if 'generated_sow' not in st.session_state:
    st.session_state.generated_sow = ""
if 'stakeholders' not in st.session_state:
    st.session_state.stakeholders = {
        "Partner": pd.DataFrame([{"Name": "Gaurav Kankaria", "Title": "Head of Analytics & ML", "Email": "gaurav.kankaria@oneture.com"}]),
        "Customer": pd.DataFrame([{"Name": "Cheten Dev", "Title": "Head of Product Design", "Email": "cheten.dev@nykaa.com"}]),
        "AWS": pd.DataFrame([{"Name": "Anubhav Sood", "Title": "AWS Account Executive", "Email": "anbhsood@amazon.com"}]),
        "Escalation": pd.DataFrame([
            {"Name": "Omkar Dhavalikar", "Title": "AI/ML Lead", "Email": "omkar.dhavalikar@oneture.com"},
            {"Name": "Gaurav Kankaria", "Title": "Head of Analytics and AIML", "Email": "gaurav.kankaria@oneture.com"}
        ])
    }

# --- SIDEBAR ---
with st.sidebar:
    st.image("https://img.icons8.com/fluency/96/artificial-intelligence.png", width=60)
    st.title("SOW Architect")
    st.caption("Enterprise POC/MVP Engine")
    with st.expander("🔑 API Key", expanded=False):
        api_key = st.text_input("Gemini API Key", type="password")
    st.divider()
    st.header("📋 1. Project Intake")
    sow_type_options = list(SOW_DIAGRAM_MAP.keys())
    selected_sow_name = st.selectbox("1.1 Scope of Work Type", sow_type_options)
    st.divider()
    industry_options = ["Retail / E-commerce", "BFSI", "Manufacturing", "Telecom", "Healthcare", "Energy / Utilities", "Logistics", "Media", "Government", "Other (specify)"]
    industry_type = st.selectbox("1.2 Industry / Domain", industry_options)
    final_industry = st.text_input("Specify Industry", placeholder="Enter industry...") if industry_type == "Other (specify)" else industry_type
    duration = st.text_input("Timeline / Duration", "4 Weeks")

# --- MAIN UI ---
st.title("🚀 GenAI Scope of Work Architect")
st.header("📸 Cover Page Branding")
brand_col1, brand_col2 = st.columns(2)
with brand_col1:
    aws_pn_logo = st.file_uploader("Top Left: AWS Partner Network Logo", type=['png','jpg','jpeg'], key="aws_pn")
    customer_logo = st.file_uploader("Slot 1: Customer Logo", type=['png','jpg','jpeg'], key="cust_logo")
with brand_col2:
    oneture_logo = st.file_uploader("Slot 2: Oneture Logo", type=['png','jpg','jpeg'], key="one_logo")
    aws_adv_logo = st.file_uploader("Slot 3: AWS Advanced Logo", type=['png','jpg','jpeg'], key="aws_adv")
    doc_date = st.date_input("Document Date", date.today())

st.divider()

# --- STEP 2: OBJECTIVES & STAKEHOLDERS ---
st.header("2. Objectives & Stakeholders")

st.subheader("🎯 2.1 Objective")
objective = st.text_area(
    "Define the core business objective:", 
    placeholder="e.g., Development of a Gen AI based WIMO Bot to demonstrate feasibility...",
    height=120
)
outcomes = st.multiselect(
    "Select success metrics:", 
    ["Reduced Response Time", "Automated SOP Mapping", "Cost Savings", "Higher Accuracy", "Metadata Richness", "Revenue Growth", "Security Compliance", "Scalability", "Integration Feasibility"],
    default=["Higher Accuracy", "Cost Savings"]
)

st.divider()

st.subheader("👥 2.2 Project Sponsor(s) / Stakeholder(s) / Project Team")
col_team1, col_team2 = st.columns(2)

with col_team1:
    st.markdown('<div class="stakeholder-header">Partner Executive Sponsor</div>', unsafe_allow_html=True)
    st.session_state.stakeholders["Partner"] = st.data_editor(st.session_state.stakeholders["Partner"], num_rows="dynamic", use_container_width=True, key="ed_partner")

    st.markdown('<div class="stakeholder-header">AWS Executive Sponsor</div>', unsafe_allow_html=True)
    st.session_state.stakeholders["AWS"] = st.data_editor(st.session_state.stakeholders["AWS"], num_rows="dynamic", use_container_width=True, key="ed_aws")

with col_team2:
    st.markdown('<div class="stakeholder-header">Customer Executive Sponsor</div>', unsafe_allow_html=True)
    st.session_state.stakeholders["Customer"] = st.data_editor(st.session_state.stakeholders["Customer"], num_rows="dynamic", use_container_width=True, key="ed_customer")

    st.markdown('<div class="stakeholder-header">Project Escalation Contacts</div>', unsafe_allow_html=True)
    st.session_state.stakeholders["Escalation"] = st.data_editor(st.session_state.stakeholders["Escalation"], num_rows="dynamic", use_container_width=True, key="ed_escalation")

# --- GENERATION ---
if st.button("✨ Generate SOW Document", type="primary", use_container_width=True):
    if not api_key:
        st.warning("⚠️ Enter a Gemini API Key in the sidebar.")
    elif not objective:
        st.error("⚠️ Business Objective is required.")
    else:
        import requests
        with st.spinner(f"Architecting {selected_sow_name}..."):
            url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash-preview-09-2025:generateContent?key={api_key}"
            
            def get_md(df):
                return df.to_markdown(index=False)

            prompt_text = f"""
            Generate a COMPLETE formal enterprise Scope of Work (SOW) for {selected_sow_name} in {final_industry}.
            
            STRICT PAGE & SECTION FLOW:
            1 TABLE OF CONTENTS (Indented sub-items)
            2 PROJECT OVERVIEW
              2.1 OBJECTIVE (Strictly 2-3 lines based on user input: {objective})
              2.2 PROJECT SPONSOR(S) / STAKEHOLDER(S) / PROJECT TEAM
                  You MUST display the following FOUR sections clearly and distinctly, each with its own heading followed by the corresponding table:
                  ### Partner Executive Sponsor
                  {get_md(st.session_state.stakeholders["Partner"])}
                  
                  ### Customer Executive Sponsor
                  {get_md(st.session_state.stakeholders["Customer"])}
                  
                  ### AWS Executive Sponsor
                  {get_md(st.session_state.stakeholders["AWS"])}
                  
                  ### Project Escalation Contacts
                  {get_md(st.session_state.stakeholders["Escalation"])}
              2.3 ASSUMPTIONS & DEPENDENCIES
              2.4 PoC Success Criteria
            3 SCOPE OF WORK – TECHNICAL PROJECT PLAN
            4 SOLUTION ARCHITECTURE / ARCHITECTURAL DIAGRAM
            6 RESOURCES & COST ESTIMATES

            CONTENT REQUIREMENTS FOR 2.4 (PoC Success Criteria):
            Strictly include these 5 outcomes:
            1. Accurate Compliance Validation: Accurate detection of compliance/non-compliance against design guidelines; identification of errors (blocking) vs warnings (quality).
            2. Structured Metadata (Tags) Extraction: Auto-generation of tags including compliance status, CTA type, Offer type, Products shown, Brands shown, and Brand ambassador presence.
            3. Ad Score Generation: Working framework (0-100) reflecting quality and compliance.
            4. Recommendations & Feedback: Clear actionable recommendations (e.g. "increase resolution") aligned with guidelines.
            5. Usability & Workflow Demonstration: Seamless end-to-end flow: Upload -> Compliance -> Summary -> Score -> Recommendations.

            CONTENT REQUIREMENTS FOR 3 (SCOPE OF WORK - TECHNICAL PROJECT PLAN):
            Strictly include these 4 phases:
            1. Infrastructure Setup: Setup AWS services (Bedrock, S3, Lambda, etc.) and gather samples/guidelines.
            2. Create Core Workflows: Banner Upload & Validation, Compliance & Tagging Flow, Issue Detection & Recommendation Flow, Ad Scoring Flow.
            3. Backend Components: Implement Compliance Engine, build Tagging Module, and store in Amazon S3.
            4. Testing and Feedback: Create PoC UI, validate accuracy against manual reviewer results, and gather stakeholder feedback.

            CONTENT RULES:
            - Section 4 must include the text: "Specifics to be discussed basis POC".
            - NO filler text or introductory sentences between headers.
            - Remove ALL markdown bolding marks (**) inside headings or body text.
            - Use plain text output only.

            INPUT DETAILS:
            - SOW Document Type: {selected_sow_name}
            - Timeline: {duration}
            
            Tone: Professional consulting. Output: Markdown only.
            """
            
            payload = {
                "contents": [{"parts": [{"text": prompt_text}]}],
                "systemInstruction": {"parts": [{"text": "You are a senior Solutions Architect. You generate detailed SOW documents. Strictly follow numbering and flow. Ensure stakeholder sections in 2.2 are distinct with their own sub-headers and tables. Sections 2.4 and 3 must be comprehensive as described. No markdown bolding."}]}
            }
            
            try:
                res = requests.post(url, json=payload)
                if res.status_code == 200:
                    st.session_state.generated_sow = res.json()['candidates'][0]['content']['parts'][0]['text']
                    st.balloons()
                else:
                    st.error(f"API Error: {res.text}")
            except Exception as e:
                st.error(f"Error: {str(e)}")

# --- STEP 3: REVIEW & EXPORT ---
if st.session_state.generated_sow:
    st.divider()
    st.header("3. Review & Export")
    tab_edit, tab_preview = st.tabs(["✍️ Document Editor", "📄 Visual Preview"])
    
    with tab_edit:
        st.session_state.generated_sow = st.text_area(
            label="Modify generated content:", 
            value=st.session_state.generated_sow, 
            height=700, 
            key="sow_editor"
        )
    
    with tab_preview:
        st.markdown(f'<div class="sow-preview">', unsafe_allow_html=True)
        header_pattern = r'(?i)(^#\s*4\s+SOLUTION ARCHITECTURE.)'
        match = re.search(header_pattern, st.session_state.generated_sow, re.MULTILINE)
        
        if match:
            start, end = match.span()
            st.markdown(st.session_state.generated_sow[:end])
            diagram_path_out = SOW_DIAGRAM_MAP.get(selected_sow_name)
            if diagram_path_out and os.path.exists(diagram_path_out):
                st.image(diagram_path_out, caption=f"{selected_sow_name} Architecture", use_container_width=True)
            st.markdown(st.session_state.generated_sow[end:])
        else:
            st.markdown(st.session_state.generated_sow)
        st.markdown('</div>', unsafe_allow_html=True)
    
    st.write("")
    
    if st.button("💾 Prepare Microsoft Word Document"):
        branding_info = {
            'sow_name': selected_sow_name,
            'aws_pn_logo_bytes': aws_pn_logo.getvalue() if aws_pn_logo else None,
            'customer_logo_bytes': customer_logo.getvalue() if customer_logo else None,
            'oneture_logo_bytes': oneture_logo.getvalue() if oneture_logo else None,
            'aws_adv_logo_bytes': aws_adv_logo.getvalue() if aws_adv_logo else None,
            'doc_date_str': doc_date.strftime("%d %B %Y")
        }
        
        docx_data = create_docx_logic(st.session_state.generated_sow, branding_info, selected_sow_name)
        
        st.download_button(
            label="📥 Download Now (.docx)", 
            data=docx_data, 
            file_name=f"SOW_{selected_sow_name.replace(' ', '_')}.docx", 
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
            use_container_width=True
        )
